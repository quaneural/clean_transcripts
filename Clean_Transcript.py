import tkinter as tk
from tkinter import filedialog
import shutil
import os
import re
import nltk
from docx import Document
from collections import defaultdict


class Application(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.pack()
        self.create_widgets()

    def create_widgets(self):
        # Create a label asking for a filename
        self.filename_label = tk.Label(self, text="Enter a filename:")
        self.filename_label.pack(side="top")

        # Create a text entry widget to enter the filename
        self.filename_entry = tk.Entry(self)
        self.filename_entry.pack(side="top")

        # Create a button to select the file and set its path
        self.select_file_button = tk.Button(self)
        self.select_file_button["text"] = "Select File"
        self.select_file_button["command"] = self.select_file
        self.select_file_button.pack(side="top")

        # Create a button to clean up the file
        self.clean_file_button = tk.Button(self)
        self.clean_file_button["text"] = "Clean up file"
        self.clean_file_button["command"] = self.clean_file
        self.clean_file_button.pack(side="top")

        # Create a button to quit the application
        self.quit = tk.Button(self, text="QUIT", fg="red",
                              command=self.master.destroy)
        self.quit.pack(side="bottom")

    def select_file(self):
        # Open a file dialog and get the selected file path
        filepath = filedialog.askopenfilename()

        # Set the filename entry widget to the selected file path
        self.filename_entry.delete(0, tk.END)
        self.filename_entry.insert(0, filepath)

    def clean_file(self):
        # Get the selected file path from the filename entry widget
        filepath = self.filename_entry.get()

        # Check if the file exists
        if not os.path.isfile(filepath):
            print("File not found!")
            return

        # Create a copy of the file with the same name and add "_cleaned.txt" to the end
        cleaned_filepath = os.path.splitext(filepath)[0] + "_cleaned.txt"

        # Download the necessary nltk data if it's not already downloaded
        nltk.download('punkt')
        nltk.download('averaged_perceptron_tagger')

        document = Document(filepath)
        text = []
        
        for data in document.paragraphs:
            text.append(data.text)

        # Combine the text into a single string and remove unnecessary content
        data = '\n'.join(text)
        data = re.sub(r'Transcription by.*?\n', '', data)
        data = re.sub(r'Speaker [0-9]*:\n', '', data)

        # Remove timestamps and other unnecessary content
        data = re.sub(r"\d{1,2}:\d{1,2}:\d{1,2}\.\d{1,3} --> \d{1,2}:\d{1,2}:\d{1,2}\.\d{1,3}", "", data)


        speakers = {name: name for line in [data.text for data in document.paragraphs]
                    for words in [line.split()]
                    for i, name in enumerate([f'{w} {words[i+1]}' for i, w in enumerate(words) 
                                               if w[0].isupper() and not w.startswith(('0:', '1:', '2:', '3:', '4:', '5:', '6:', '7:', '8:', '9:')) 
                                               if i < len(words) - 1]) 
                    if all(c.isalpha() or c.isspace() for c in name) 
                    and any(tag == 'NNP' for word, tag in nltk.pos_tag(nltk.word_tokenize(name)))}


        # Print the extracted full names for verification
        speakers.update({"Tom O'Shea": "Tom O'Shea"})
        print(speakers)

        # Split the text into lines and group them by speaker name
        groups = []
        current_speaker = None
        for line in data.split('\n'):
            if line in speakers:
                current_speaker = speakers[line]
            elif current_speaker is not None:
                # If the current line is empty, skip it
                if line.strip() == '':
                    continue
                # If the previous line was also from the same speaker, append the current line to it
                if len(groups) > 0 and groups[-1][0] == current_speaker:
                    groups[-1][1].append(line)
                else:
                    groups.append((current_speaker, [line]))

        # Write the grouped text to a new file, in the order of the conversation
        with open(cleaned_filepath, 'w') as file:
            for speaker, lines in groups:
                # Combine the lines into a single string
                message = ' '.join(lines)

                # Remove the timestamp from the beginning of the message
                message = ' '.join(message.split()[1:])

                file.write(f'{speaker}:\n{message}\n\n')
            file.write('\n')

        print("File cleaned up successfully!")


root = tk.Tk()
app = Application(master=root)
app.mainloop()
